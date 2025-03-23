import os
import google.generativeai as genai
from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from pydantic import BaseModel
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import io
from fastapi.responses import StreamingResponse

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise RuntimeError("❌ Google API key is missing. Set 'GEMINI_API_KEY' in .env file.")

# Configure Google Generative AI
genai.configure(api_key=GEMINI_API_KEY)

# Initialize FastAPI App
app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define Pydantic Model
class ChatRequest(BaseModel):
    user_message: str

# AI Model Configuration
safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
]

generation_config = {
    "temperature": 0.7,
    "top_p": 0.9,
    "max_output_tokens": 1000,
}

model = genai.GenerativeModel(
    model_name="gemini-2.0-pro-exp-02-05",
    safety_settings=safety_settings,
    generation_config=generation_config,
)

DISCLAIMER = "\n\n*Disclaimer: This AI-generated response is for informational purposes only and should not be considered legal advice.*"

LEGAL_PRE_PROMPT = """
You are an AI legal assistant with expertise in contract drafting and legal advisory. 
Your response must be clear, legally precise, and formatted professionally.

Task:
1. Generate a well-structured **legal contract** based on the provided clauses.
2. Provide **legal advice** related to the clauses, including risks, improvements, and missing terms.

Make sure to use formal legal language and be concise.
"""

@app.get("/")
def root():
    return {"message": "Welcome to Pocket Advocate API!"}

@app.post("/chat/uploadImage")
async def upload_image(user_message: str = Form(""), image: UploadFile = File(None)):
    try:
        image_data = await image.read() if image else None
        full_prompt = f"{LEGAL_PRE_PROMPT}\nUser: {user_message}\n" + ("Image attached." if image_data else "")
        
        response = model.generate_content(full_prompt)
        bot_response = response.text if hasattr(response, "text") else "No response generated."

        return {"user_message": user_message, "bot_response": bot_response + DISCLAIMER, "image_uploaded": bool(image_data)}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image upload error: {str(e)}")

@app.post("/chat/chatWithBot")
def chat_with_bot(chat_request: ChatRequest):
    try:
        full_prompt = f"{LEGAL_PRE_PROMPT}\nUser: {chat_request.user_message}"
        response = model.generate_content(full_prompt)
        bot_response = response.text if hasattr(response, "text") else "Sorry, I couldn't generate a response."

        return {"user_message": chat_request.user_message, "bot_response": bot_response + DISCLAIMER}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chatbot error: {str(e)}")

@app.post("/chat/summarizeDocument")
async def summarize_document(file: UploadFile = File(...)):
    try:
        text = (await file.read()).decode("utf-8")
        summary_prompt = f"Summarize this legal document concisely:\n{text}"

        response = model.generate_content(summary_prompt)
        bot_response = response.text if hasattr(response, "text") else "Sorry, could not generate a summary."

        return {"summary": bot_response}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document summarization error: {str(e)}")

@app.post("/chat/contractQualityCheck")
async def contract_quality_check(file: UploadFile = File(...)):
    try:
        text = (await file.read()).decode("utf-8")
        analysis_prompt = f"Analyze this contract for missing clauses, legal risks, and areas of improvement:\n{text}"

        response = model.generate_content(analysis_prompt)
        bot_response = response.text if hasattr(response, "text") else "Sorry, could not analyze the contract."

        return {"contract_analysis": bot_response}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Contract analysis error: {str(e)}")

@app.post("/chat/downloadContract")
async def download_contract(chat_request: ChatRequest):
    try:
        full_prompt = f"{LEGAL_PRE_PROMPT}\nUser: {chat_request.user_message}"
        response = model.generate_content(full_prompt)
        bot_response = response.text if hasattr(response, "text") else "Sorry, could not generate the contract."

        # Create DOCX file
        doc = Document()
        doc.add_heading("AI-Generated Legal Contract", level=1)
        doc.add_paragraph(bot_response + DISCLAIMER)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=contract.docx"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Contract generation error: {str(e)}")
